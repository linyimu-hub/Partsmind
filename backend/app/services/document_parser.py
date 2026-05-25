"""
app/services/document_parser.py
────────────────────────────────
Parses PDF and DOCX files into clean text chunks.

Chunking strategy:
  - Target: 300 tokens per chunk (≈ 1200 chars)
  - Overlap: 50 tokens (≈ 200 chars) between adjacent chunks
  - Why overlap? Ensures context isn't lost at chunk boundaries.
    A question about content spanning two chunks can still be answered.
  - Split preference: paragraph > sentence > character
    We try to split at natural language boundaries first.

Each chunk carries metadata (page, section) for citation display.
"""

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from app.core.logging import get_logger

logger = get_logger(__name__)

# Target chunk size in characters (≈ 300 tokens for English, ~400 for Chinese)
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 200
MIN_CHUNK_SIZE = 100   # discard chunks smaller than this (usually headers/noise)


@dataclass
class ParsedChunk:
    content: str
    chunk_index: int
    metadata: dict[str, Any] = field(default_factory=dict)
    # metadata example: {"page": 3, "section": "Return Policy", "char_start": 1240}


@dataclass
class ParsedDocument:
    full_text: str
    chunks: list[ParsedChunk]
    page_count: int
    metadata: dict[str, Any] = field(default_factory=dict)


# ── PDF Parser ────────────────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes) -> ParsedDocument:
    """
    Extract text from PDF page by page.
    Tracks page numbers for citation display.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError("pypdf not installed. Run: pip install pypdf")

    reader = PdfReader(io.BytesIO(file_bytes))
    page_count = len(reader.pages)

    # Extract text with page tracking
    page_texts: list[tuple[int, str]] = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = _clean_text(text)
        if text.strip():
            page_texts.append((page_num, text))

    full_text = "\n\n".join(text for _, text in page_texts)

    logger.info("parser.pdf_extracted", pages=page_count, chars=len(full_text))

    # Chunk with page metadata
    chunks = _chunk_with_page_metadata(page_texts)

    return ParsedDocument(
        full_text=full_text,
        chunks=chunks,
        page_count=page_count,
        metadata={"format": "pdf", "pages": page_count},
    )


# ── DOCX Parser ───────────────────────────────────────────────────────────────

def parse_docx(file_bytes: bytes) -> ParsedDocument:
    """
    Extract text from DOCX, preserving heading structure AND table content.
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document(io.BytesIO(file_bytes))

    paragraphs: list[dict[str, Any]] = []
    current_section = "Introduction"

    # 提取段落 + 表格，按文档顺序遍历
    from docx.oxml.ns import qn
    body = doc.element.body
    
    for child in body.iterchildren():
        # 段落
        if child.tag == qn('w:p'):
            text = ''.join(t.text or '' for t in child.iter(qn('w:t'))).strip()
            if not text:
                continue
            # 检查是否是标题
            pPr = child.find(qn('w:pPr'))
            is_heading = False
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_val = pStyle.get(qn('w:val'), '')
                    if 'Heading' in style_val or 'heading' in style_val:
                        is_heading = True
                        current_section = text
            paragraphs.append({
                "text": text,
                "section": current_section,
                "is_heading": is_heading,
            })
        
        # 表格
        elif child.tag == qn('w:tbl'):
            table_text_parts = []
            for row in child.iter(qn('w:tr')):
                row_cells = []
                for cell in row.iter(qn('w:tc')):
                    cell_text = ''.join(t.text or '' for t in cell.iter(qn('w:t'))).strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    table_text_parts.append(" | ".join(row_cells))
            if table_text_parts:
                paragraphs.append({
                    "text": "\n".join(table_text_parts),
                    "section": current_section,
                    "is_heading": False,
                })

    full_text = "\n\n".join(p["text"] for p in paragraphs)
    logger.info("parser.docx_extracted", paragraphs=len(paragraphs), chars=len(full_text))

    chunks = _chunk_with_section_metadata(paragraphs)

    return ParsedDocument(
        full_text=full_text,
        chunks=chunks,
        page_count=0,
        metadata={"format": "docx", "sections": list({p["section"] for p in paragraphs})},
    )# ── Dispatcher ────────────────────────────────────────────────────────────────

def parse_document(file_bytes: bytes, file_type: str) -> ParsedDocument:
    """Route to the correct parser based on file type."""
    if file_type == "pdf":
        return parse_pdf(file_bytes)
    elif file_type in ("docx", "doc"):
        return parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


# ── Chunking helpers ──────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Remove PDF artifacts: multiple spaces, form feeds, etc."""
    text = re.sub(r'\f', '\n', text)           # form feed → newline
    text = re.sub(r' {3,}', '  ', text)        # 3+ spaces → 2 spaces
    text = re.sub(r'\n{4,}', '\n\n\n', text)   # 4+ newlines → 3
    return text.strip()


def _split_into_segments(text: str) -> list[str]:
    """
    Split text at paragraph boundaries first, then sentences.
    This produces more semantically coherent chunks than
    naive character splitting.
    """
    # Try paragraph split first
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    if not paragraphs:
        paragraphs = [text]

    segments: list[str] = []
    for para in paragraphs:
        if len(para) <= CHUNK_SIZE:
            segments.append(para)
        else:
            # Split long paragraph at sentence boundaries
            sentences = re.split(r'(?<=[。！？.!?])\s+', para)
            segments.extend(s for s in sentences if s.strip())

    return segments


def _build_chunks_from_segments(
    segments: list[str],
    metadata_fn,   # callable(char_position) → dict
) -> list[ParsedChunk]:
    """
    Merge segments into chunks of ~CHUNK_SIZE chars with CHUNK_OVERLAP.
    """
    chunks: list[ParsedChunk] = []
    current_text = ""
    current_start = 0
    chunk_index = 0
    char_pos = 0

    for segment in segments:
        # If adding this segment would exceed chunk size, finalize current chunk
        if current_text and len(current_text) + len(segment) + 1 > CHUNK_SIZE:
            if len(current_text) >= MIN_CHUNK_SIZE:
                chunks.append(ParsedChunk(
                    content=current_text.strip(),
                    chunk_index=chunk_index,
                    metadata=metadata_fn(current_start),
                ))
                chunk_index += 1

            # Start new chunk with overlap from end of current
            overlap_text = current_text[-CHUNK_OVERLAP:] if len(current_text) > CHUNK_OVERLAP else current_text
            current_text = overlap_text + " " + segment
            current_start = char_pos - len(overlap_text)
        else:
            current_text = (current_text + " " + segment).strip() if current_text else segment

        char_pos += len(segment) + 1

    # Don't forget the last chunk
    if current_text.strip() and len(current_text) >= MIN_CHUNK_SIZE:
        chunks.append(ParsedChunk(
            content=current_text.strip(),
            chunk_index=chunk_index,
            metadata=metadata_fn(current_start),
        ))

    return chunks


def _chunk_with_page_metadata(
    page_texts: list[tuple[int, str]]
) -> list[ParsedChunk]:
    """Chunk PDF text, tracking which page each chunk came from."""
    # Build a map: char_offset → page_number
    page_map: list[tuple[int, int]] = []
    char_pos = 0
    for page_num, text in page_texts:
        page_map.append((char_pos, page_num))
        char_pos += len(text) + 2

    def get_page(char_start: int) -> int:
        page = 1
        for offset, num in page_map:
            if char_start >= offset:
                page = num
        return page

    full_text = "\n\n".join(text for _, text in page_texts)
    segments = _split_into_segments(full_text)

    return _build_chunks_from_segments(
        segments,
        lambda pos: {"page": get_page(pos), "char_start": pos},
    )


def _chunk_with_section_metadata(
    paragraphs: list[dict[str, Any]]
) -> list[ParsedChunk]:
    """Chunk DOCX text, tracking section headings."""
    segments_with_meta: list[tuple[str, str]] = [
        (p["text"], p["section"]) for p in paragraphs
    ]

    chunks: list[ParsedChunk] = []
    current_text = ""
    current_section = ""
    chunk_index = 0

    for text, section in segments_with_meta:
        if current_text and len(current_text) + len(text) + 1 > CHUNK_SIZE:
            if len(current_text) >= MIN_CHUNK_SIZE:
                chunks.append(ParsedChunk(
                    content=current_text.strip(),
                    chunk_index=chunk_index,
                    metadata={"section": current_section},
                ))
                chunk_index += 1
            overlap = current_text[-CHUNK_OVERLAP:]
            current_text = overlap + " " + text
        else:
            current_text = (current_text + " " + text).strip() if current_text else text
            current_section = section

    if current_text.strip() and len(current_text) >= MIN_CHUNK_SIZE:
        chunks.append(ParsedChunk(
            content=current_text.strip(),
            chunk_index=chunk_index,
            metadata={"section": current_section},
        ))

    return chunks
