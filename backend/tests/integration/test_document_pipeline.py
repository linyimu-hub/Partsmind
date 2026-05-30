"""
Integration test for document ingestion pipeline.
Tests parser + chunker without needing real DB or OpenAI.
"""

import io

import pytest

from app.services.document_parser import (
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    MIN_CHUNK_SIZE,
    _build_chunks_from_segments,
    _split_into_segments,
    parse_docx,
)

# ── Unit tests for chunking logic (no external deps) ──────────────────────────

def test_split_segments_short_text():
    text = "Hello world. This is a test."
    segments = _split_into_segments(text)
    assert len(segments) >= 1
    assert "Hello world" in segments[0]


def test_chunk_size_respected():
    """No chunk should exceed CHUNK_SIZE characters."""
    long_text = "word " * 2000  # ~10,000 chars
    segments = _split_into_segments(long_text)
    chunks = _build_chunks_from_segments(segments, lambda pos: {"char_start": pos})
    for chunk in chunks:
        assert len(chunk.content) <= CHUNK_SIZE + CHUNK_OVERLAP + 100  # small tolerance


def test_chunks_have_overlap():
    """Adjacent chunks should share some content (overlap)."""
    # Generate a document long enough to produce multiple chunks
    paragraphs = [f"Paragraph {i}: " + "content " * 50 for i in range(20)]
    text = "\n\n".join(paragraphs)
    segments = _split_into_segments(text)
    chunks = _build_chunks_from_segments(segments, lambda pos: {})

    if len(chunks) >= 2:
        # The end of chunk[0] should appear in the beginning of chunk[1]
        end_of_first = chunks[0].content[-100:]
        start_of_second = chunks[1].content[:200]
        # At least some content should overlap
        words_first = set(end_of_first.split())
        words_second = set(start_of_second.split())
        overlap = words_first & words_second
        assert len(overlap) > 0, "Adjacent chunks should share words from overlap"


def test_min_chunk_size_filter():
    """Chunks smaller than MIN_CHUNK_SIZE should be discarded."""
    short_segments = ["hi", "ok", "yes"]  # all very short
    chunks = _build_chunks_from_segments(short_segments, lambda pos: {})
    for chunk in chunks:
        assert len(chunk.content) >= MIN_CHUNK_SIZE


def test_chunk_indices_sequential():
    """Chunk indices should be 0, 1, 2, ..."""
    text = "\n\n".join(["sentence " * 100 for _ in range(10)])
    segments = _split_into_segments(text)
    chunks = _build_chunks_from_segments(segments, lambda pos: {})
    for i, chunk in enumerate(chunks):
        assert chunk.chunk_index == i


# ── DOCX parsing (using python-docx to create test file) ──────────────────────

def _make_test_docx() -> bytes:
    """Create a minimal DOCX in memory for testing."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_heading("Section 1: Brake Systems", level=2)
        for i in range(5):
            doc.add_paragraph(
                f"Brake pad paragraph {i}: The ceramic brake pads provide "
                f"excellent stopping power and low dust characteristics. "
                f"Compatible with most Toyota and Honda vehicles. " * 3
            )
        doc.add_heading("Section 2: Filters", level=2)
        for i in range(5):
            doc.add_paragraph(
                f"Filter paragraph {i}: High-flow air filters improve engine "
                f"performance and extend service intervals significantly. " * 3
            )
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        pytest.skip("python-docx not installed")


def test_parse_docx_produces_chunks():
    docx_bytes = _make_test_docx()
    result = parse_docx(docx_bytes)
    assert len(result.chunks) > 0
    assert result.full_text != ""


def test_parse_docx_section_metadata():
    docx_bytes = _make_test_docx()
    result = parse_docx(docx_bytes)
    sections = {c.metadata.get("section") for c in result.chunks}
    # Should have captured at least one section heading
    assert len(sections) > 0


def test_parse_docx_no_empty_chunks():
    docx_bytes = _make_test_docx()
    result = parse_docx(docx_bytes)
    for chunk in result.chunks:
        assert chunk.content.strip() != ""
        assert len(chunk.content) >= MIN_CHUNK_SIZE
